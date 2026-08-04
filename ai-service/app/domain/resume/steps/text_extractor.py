"""Extract plain text from PDF and DOCX resume documents."""

import io
import logging

from app.domain.resume.exceptions import ResumeValidationError
from app.domain.resume.steps.file_validator import DOCX_MIME, PDF_MIME

logger = logging.getLogger(__name__)

MAX_PDF_PAGES = 10
OCR_TEXT_THRESHOLD = 100


def _extract_pdf_text(file_bytes: bytes) -> tuple[str, int]:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(file_bytes))
    page_count = len(reader.pages)
    pages_text = [
        text for page in reader.pages[:MAX_PDF_PAGES] if (text := page.extract_text())
    ]
    return "\n\n".join(pages_text), page_count


def _extract_pdf_with_ocr(file_bytes: bytes) -> str:
    try:
        from pdf2image import convert_from_bytes
        from pytesseract import image_to_string
    except ImportError as exc:
        raise ResumeValidationError("OCR dependencies are not installed") from exc

    try:
        images = convert_from_bytes(
            file_bytes,
            dpi=250,
            first_page=1,
            last_page=MAX_PDF_PAGES,
            thread_count=1,
            timeout=60,
        )
        return "\n\n".join(
            image_to_string(image, lang="vie+eng", timeout=30) for image in images
        )
    except Exception as exc:
        raise ResumeValidationError(f"PDF OCR failed: {exc}") from exc


def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        text, page_count = _extract_pdf_text(file_bytes)
    except Exception as exc:
        raise ResumeValidationError(f"PDF text extraction failed: {exc}") from exc

    if page_count > MAX_PDF_PAGES:
        logger.warning(
            "PDF has %d pages; only the first %d are processed",
            page_count,
            MAX_PDF_PAGES,
        )

    if len(text.strip()) < OCR_TEXT_THRESHOLD:
        logger.info("PDF contains little embedded text; attempting OCR")
        try:
            ocr_text = _extract_pdf_with_ocr(file_bytes)
            if ocr_text.strip():
                text = ocr_text
        except ResumeValidationError:
            if not text.strip():
                raise
            logger.warning("OCR failed; continuing with limited embedded text")

    logger.info("Extracted %d characters from PDF", len(text))
    return text


def extract_text_from_docx(file_bytes: bytes) -> str:
    from docx import Document

    try:
        document = Document(io.BytesIO(file_bytes))
        parts: list[str] = []

        def append(value: str) -> None:
            cleaned = value.strip()
            if cleaned:
                parts.append(cleaned)

        for paragraph in document.paragraphs:
            append(paragraph.text)

        for table in document.tables:
            for row in table.rows:
                append(" | ".join(cell.text.strip() for cell in row.cells))

        for section in document.sections:
            for container in (section.header, section.footer):
                for paragraph in container.paragraphs:
                    append(paragraph.text)

        # Text boxes are not exposed by python-docx's high-level API.
        for node in document.element.xpath(".//w:txbxContent//w:t"):
            append(node.text or "")

        result = "\n".join(parts)
        logger.info("Extracted %d characters from DOCX", len(result))
        return result
    except Exception as exc:
        raise ResumeValidationError(f"DOCX text extraction failed: {exc}") from exc


def extract_text(file_bytes: bytes, mime_type: str) -> str:
    if mime_type == PDF_MIME:
        return extract_text_from_pdf(file_bytes)
    if mime_type == DOCX_MIME:
        return extract_text_from_docx(file_bytes)
    raise ResumeValidationError(f"Unsupported MIME type: {mime_type}")
