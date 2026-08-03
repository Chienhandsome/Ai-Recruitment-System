import io

from docx import Document

from app.domain.resume.steps import text_extractor


def test_docx_extractor_reads_paragraph_table_header_and_footer():
    document = Document()
    document.add_paragraph("Backend engineer")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Python"
    table.cell(0, 1).text = "FastAPI"
    document.sections[0].header.paragraphs[0].text = "Candidate CV"
    document.sections[0].footer.paragraphs[0].text = "Page footer"
    stream = io.BytesIO()
    document.save(stream)

    result = text_extractor.extract_text_from_docx(stream.getvalue())

    assert "Backend engineer" in result
    assert "Python | FastAPI" in result
    assert "Candidate CV" in result
    assert "Page footer" in result


def test_pdf_extractor_uses_ocr_when_embedded_text_is_short(monkeypatch):
    monkeypatch.setattr(
        text_extractor,
        "_extract_pdf_text",
        lambda _: ("too short", 1),
    )
    monkeypatch.setattr(
        text_extractor,
        "_extract_pdf_with_ocr",
        lambda _: "OCR extracted resume text",
    )

    assert text_extractor.extract_text_from_pdf(b"pdf") == (
        "OCR extracted resume text"
    )
